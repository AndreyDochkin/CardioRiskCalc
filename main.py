import tkinter as tk
from tkinter import messagebox
from tkinter import simpledialog
from docx import Document
import os

def calculate_points_SA(SA):
    if SA <= 19: return 1
    elif SA <= 3: return 2
    else: return 3

def calculate_points_IAAK(IAAK):
    if IAAK < 2: return 1
    elif IAAK <= 5: return 2
    else: return 3

def calculate_points_VT(VT):
    if VT <= 8: return 1
    elif VT <= 10: return 2
    else: return 3

def calculate_points_FG(FG):
    if FG == "2С19*17": return 1
    elif FG == "2С19*1": return 2
    else: return 3

def calculate_points_IAADF5(IAADF5):
    if IAADF5 < 15: return 1
    elif IAADF5 <= 60: return 2
    else: return 3

def calculate_risk():
    SA = float(entry_SA.get())
    IAAK = float(entry_IAAK.get())
    VT = float(entry_VT.get())
    FG = combo_FG.get()
    IAADF5 = float(entry_IAADF5.get())
    
    points = 0
    points += calculate_points_SA(SA)
    points += calculate_points_IAAK(IAAK)
    points += calculate_points_VT(VT)
    points += calculate_points_FG(FG)
    points += calculate_points_IAADF5(IAADF5)
    
    Krs = points / 5
    
    generate_doc(Krs)

def generate_doc(Krs):
    document = Document()
    document.add_heading('Результаты расчета Крс', 0)
    
    if Krs < 2:
        recommendation = "Целевая гипоагрегация достигнута. Коррекция терапии не требуется. Риск повторных сосудистых событий низкий."
    elif Krs <= 2.8:
        recommendation = "Целевая гипоагрегация не достигнута! Требуется коррекция терапии. Рекомендован переход на ДАТТ с применением комбинации препаратов: Ацетилсалициловая кислота 75 мг + Тикагрелор 90 мг х 2 раза в день. Риск повторных сосудистых событий высокий."
    else:
        recommendation = "Целевая гипоагрегация не достигнута! Требуется срочная коррекция терапии. Рекомендован переход на ДАТТ с применением комбинации препаратов: Ацетилсалициловая кислота 75 мг + Тикагрелор 90 мг х 2 раза в день. Риск повторных сосудистых событий крайне высокий. Повторное исследование в динамике."
    
    document.add_paragraph(f'Крс: {Krs:.1f}')
    document.add_paragraph('Рекомендации:')
    document.add_paragraph(recommendation)
    
    filename = 'Результаты_расчета.docx'
    document.save(filename)
    messagebox.showinfo("Готово", f"Результаты сохранены в файле {filename}")
    os.startfile(filename)

root = tk.Tk()
root.title("CardioRiskCalc")

tk.Label(root, text="Спонтанная агрегация (СА):").grid(row=0)
entry_SA = tk.Entry(root)
entry_SA.grid(row=0, column=1)

tk.Label(root, text="ИАак:").grid(row=1)
entry_IAAK = tk.Entry(root)
entry_IAAK.grid(row=1, column=1)

tk.Label(root, text="ВТ:").grid(row=2)
entry_VT = tk.Entry(root)
entry_VT.grid(row=2, column=1)

tk.Label(root, text="ФГ:").grid(row=3)
combo_FG = tk.StringVar(root)
combo_FG.set("2С19*17") # default value
tk.OptionMenu(root, combo_FG, "2С19*17", "2С19*1", "2С19*2/2С19*3").grid(row=3, column=1)

tk.Label(root, text="ИАадф5:").grid(row=4)
entry_IAADF5 = tk.Entry(root)
entry_IAADF5.grid(row=4, column=1)

tk.Button(root, text='Рассчитать', command=calculate_risk).grid(row=5, column=1, pady=4)

root.mainloop()
