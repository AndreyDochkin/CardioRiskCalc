from docx import Document
from docx.shared import Inches
import os
import sys
from PyQt5 import QtWidgets, uic
from PyQt5.QtWidgets import QApplication
from PyQt5.QtGui import QIcon

MainWindow , Form = uic.loadUiType("gui.ui")

class Ui(QtWidgets.QMainWindow, MainWindow):
    def __init__(self):
        super(Ui, self).__init__()
        self.setupUi(self)
        self.setWindowIcon(QIcon('app.ico'))
        self.pushButton.clicked.connect(self.generate_doc)
        self.lineEdit.setText('Иванов Иван')
        self.show()

    def generate_doc(self):
        # print('btn push')
        list = [self.checkBox.checkState(),
                self.checkBox_2.checkState(),
                self.checkBox_3.checkState(),
                self.checkBox_4.checkState(),
                self.checkBox_5.checkState()]

        # print(list)
        foo = list.count(2)
        # print(foo)

        document = Document()
        document.add_heading('Рекомендации по лечению', 1)
        document.add_heading('Вывод такой:', level=1)

        if foo >= 0 and foo <= 2:
            document.add_paragraph('от 0 до 2 баллов')
        elif foo > 2 and foo <= 4:
            document.add_paragraph('больше 2 баллов меньше 4')
        elif foo > 4:
            document.add_paragraph('больше 4 баллов')

        self.checkBox.setCheckState(0)
        self.checkBox_2.setCheckState(0)
        self.checkBox_3.setCheckState(0)
        self.checkBox_4.setCheckState(0)
        self.checkBox_5.setCheckState(0)

        text_name = self.lineEdit.text()
        document.save(text_name + '.docx')
        os.startfile(text_name + '.docx')


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    w = Ui()
    sys.exit(app.exec_())

# import sys
# import PyQt5
# from PyQt5.QtWidgets import QApplication
# from PyQt5.QtWidgets import QWidget, QLabel ,QDialog
# # from PyQt5 import uic
# from PyQt5 import QtCore, QtGui, QtWidgets
# import test
# from test import Ui_Form
#
#
# # class Widget(QWidget):
# #
# #     def __init__(self):
# #         super().__init__()
# #         uic.loadUi("test.ui", self)
# #         self.label.setText("NewText")
#
# class ExampleApp(QtWidgets.QDialog, Ui_Form):
#     def __init__(self):
#         # Это здесь нужно для доступа к переменным, методам
#         # и т.д. в файле design.py
#         super().__init__()
#         self.setupUi(self)  # Это нужно для инициализации нашего дизайна
#
#
# def main():
#     app = QtWidgets.QApplication(sys.argv)  # Новый экземпляр QApplication
#     window = ExampleApp()  # Создаём объект класса ExampleApp
#     window.show()  # Показываем окно
#     app.exec_()  # и запускаем приложение
#
#
# if __name__ == '__main__':  # Если мы запускаем файл напрямую, а не импортируем
#     main()  # то запускаем функцию main()
